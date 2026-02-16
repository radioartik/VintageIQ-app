import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = None

    for line in lines:
        line = line.strip()
        
        # Header 1
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header 2
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        # Header 3
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        # Lists
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Table rows
        elif line.startswith('|') and not line.startswith('|---'):
            parts = [p.strip() for p in line.split('|')[1:-1]]
            if not current_table:
                current_table = doc.add_table(rows=0, cols=len(parts))
                current_table.style = 'Table Grid'
            
            row_cells = current_table.add_row().cells
            for i, part in enumerate(parts):
                row_cells[i].text = part
        # Reset table
        elif not line.startswith('|'):
            current_table = None
            if line:
                # Basic text formatting (remove bold symbols)
                clean_line = line.replace('**', '')
                doc.add_paragraph(clean_line)

    doc.save(docx_path)

if __name__ == "__main__":
    md_to_docx('PRD.md', 'PRD.docx')
